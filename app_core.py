import os
import sys
import pandas as pd
import tkinter as tk
from tkinter import messagebox, ttk, simpledialog, filedialog
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

# =====================================================================
# 1. MOTORES PREMIUM DE EXPORTACIÓN DE ACUSACIÓN (ROI / ROR)
# =====================================================================
def crear_word_premium(nombre_archivo, tipo_reporte, datos):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    p_conf = doc.add_paragraph()
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_conf.add_run("DOCUMENTO CONFIDENCIAL DE AUDITORÍA").font.size = Pt(8.5)
    
    titulo_h = doc.add_heading(f'REPORTE DE OPERACIÓN {tipo_reporte} ({datos["siglas"]})', level=0)
    titulo_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in titulo_h.runs:
        run.font.color.rgb = RGBColor(27, 94, 32)
        run.font.size = Pt(16)
    
    p_exp = doc.add_paragraph()
    p_exp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_exp.add_run(f"Expediente: {datos['expediente']}   |   Periodo Auditado: {datos['ventana_tiempo']}")
    
    doc.add_heading('1. INSTITUCIÓN REPORTANTE', level=1)
    t_inst = doc.add_table(rows=2, cols=2)
    
    row0 = t_inst.rows[0].cells
    row0[0].text = "Institución: Mitre S.A.P.I. de C.V."
    row0[1].text = "Tipo de Entidad: IFPE"
    
    row1 = t_inst.rows[1].cells
    row1[0].text = f"Oficial / Analista: {datos['analista']}"
    row1[1].text = f"Fecha de Auditoría: {datos['fecha_auditoria']}"
    
    doc.add_heading('2. DATOS DEL CLIENTE INVESTIGADO', level=1)
    tabla_cli = doc.add_table(rows=1, cols=2)
    tabla_cli.style = 'Table Grid'
    hdr = tabla_cli.rows[0].cells
    hdr[0].text = "Campo"
    hdr[1].text = "Detalle"
    
    for campo, valor in datos['cliente_info']:
        f = tabla_cli.add_row().cells
        f[0].text = str(campo)
        f[1].text = str(valor)
        
    doc.add_heading('3. DESCRIPCIÓN DETALLADA DEL HALLAZGO', level=1)
    doc.add_paragraph(datos['resumen'])
    
    doc.add_heading('3.2 Cronología de Movimientos Involucrados', level=2)
    tabla_movs = doc.add_table(rows=1, cols=3)
    tabla_movs.style = 'Table Grid'
    hdr2 = tabla_movs.rows[0].cells
    hdr2[0].text = "Fecha / Hora"
    hdr2[1].text = "Mecanismo"
    hdr2[2].text = "Monto (MXN)"
    
    for mov in datos['movimientos_tabla']:
        f2 = tabla_movs.add_row().cells
        f2[0].text = str(mov[0])
        f2[1].text = str(mov[1])
        f2[2].text = str(mov[2])
    
    doc.add_heading('3.3 Señales de Alerta', level=2)
    for alerta in datos['alertas']:
        doc.add_paragraph(f"▲ {alerta}")

    doc.add_heading('4. CONCLUSIONES DE DEBIDA DILIGENCIA JURÍDICA', level=1)
    doc.add_paragraph(datos['entrevista'])
    doc.add_paragraph(datos['marco_legal'])
    
    doc.add_paragraph(f"\n_____________________________\nFirma Autorizada Analista: {datos['analista']}")

    # 🍏 RUTA UNIVERSAL MAC: Detecta la carpeta exacta donde guardaste tu código
    carpeta_destino = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Reportes")
    if not os.path.exists(carpeta_destino): 
        os.makedirs(carpeta_destino)
    doc.save(os.path.join(carpeta_destino, nombre_archivo))


def crear_pdf_premium(nombre_archivo, tipo_reporte, datos):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_margins(20, 20, 20)
    
    pdf.set_font('helvetica', 'I', 8)
    pdf.set_text_color(120, 120, 120)
    pdf.cell(0, 5, "DOCUMENTO CONFIDENCIAL DE AUDITORÍA", 0, 1, 'C')
    pdf.ln(2)
    
    pdf.set_font('helvetica', 'B', 15)
    pdf.set_text_color(27, 94, 32)
    pdf.cell(0, 10, f"REPORTE DE OPERACIÓN {tipo_reporte} ({datos['siglas']})", 0, 1, 'C')
    
    pdf.set_font('helvetica', '', 9)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 5, f"Expediente: {datos['expediente']} | Periodo: {datos['ventana_tiempo']}", 'B', 1, 'C')
    pdf.ln(5)
    
    pdf.set_font('helvetica', 'B', 11)
    pdf.set_text_color(27, 94, 32)
    pdf.cell(0, 6, "1. INSTITUCIÓN REPORTANTE", 0, 1, 'L')
    pdf.set_font('helvetica', '', 9)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(90, 5, "Institución: Mitre S.A.P.I. de C.V.", 0, 0, 'L')
    pdf.cell(90, 5, "Tipo de Entidad: IFPE", 0, 1, 'L')
    pdf.cell(90, 5, f"Oficial / Analista: {datos['analista']}", 0, 0, 'L')
    pdf.cell(90, 5, f"Fecha: {datos['fecha_auditoria']}", 0, 1, 'L')
    pdf.ln(4)
    
    pdf.set_font('helvetica', 'B', 11)
    pdf.set_text_color(27, 94, 32)
    pdf.cell(0, 6, "2. DATOS DEL CLIENTE INVESTIGADO", 0, 1, 'L')
    pdf.set_font('helvetica', '', 9)
    pdf.set_text_color(0, 0, 0)
    for campo, valor in datos['cliente_info']:
        pdf.cell(60, 5, f" {campo}:", 1, 0, 'L')
        pdf.cell(110, 5, f" {valor}", 1, 1, 'L')
    pdf.ln(4)
    
    pdf.set_font('helvetica', 'B', 11)
    pdf.set_text_color(27, 94, 32)
    pdf.cell(0, 6, "3. DESCRIPCIÓN DETALLADA DEL HALLAZGO", 0, 1, 'L')
    pdf.set_font('helvetica', '', 9)
    pdf.set_text_color(0, 0, 0)
    resumen_l = str(datos['resumen']).replace('•', '-')
    pdf.multi_cell(0, 5, resumen_l)
    pdf.ln(3)
    
    pdf.set_font('helvetica', 'B', 10)
    pdf.cell(50, 5, " Fecha / Hora", 1, 0, 'C')
    pdf.cell(50, 5, " Mecanismo", 1, 0, 'C')
    pdf.cell(70, 5, " Monto (MXN)", 1, 1, 'C')
    pdf.set_font('helvetica', '', 10)
    for mov in datos['movimientos_tabla']:
        pdf.cell(50, 5, f" {mov[0]}", 1, 0, 'C')
        pdf.cell(50, 5, f" {mov[1]}", 1, 0, 'C')
        pdf.cell(70, 5, f" {mov[2]}", 1, 1, 'C')
    pdf.ln(4)
    
    pdf.set_font('helvetica', 'B', 11)
    pdf.set_text_color(27, 94, 32)
    pdf.cell(0, 6, "4. CONCLUSIONES DE DEBIDA DILIGENCIA JURÍDICA", 0, 1, 'L')
    pdf.set_font('helvetica', '', 9)
    pdf.set_text_color(0, 0, 0)
    entrevista_l = str(datos['entrevista']).replace('•', '-')
    pdf.multi_cell(0, 5, entrevista_l)
    pdf.ln(2)
    pdf.multi_cell(0, 5, datos['marco_legal'])
    pdf.ln(10)
    
    pdf.cell(0, 5, '___________________________', 0, 1, 'L')
    pdf.cell(0, 5, f"Firma Autorizada Analista: {datos['analista']}", 0, 1, 'L')

    # 🍏 RUTA UNIVERSAL MAC: Detecta la carpeta exacta donde guardaste tu código
    carpeta_destino = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Reportes")
    if not os.path.exists(carpeta_destino): 
        os.makedirs(carpeta_destino)
    pdf.output(os.path.join(carpeta_destino, nombre_archivo))


# =====================================================================
# 3. EL ESCRIBANO DE LA BITÁCORA ANUALIZADA (EXCEL)
# =====================================================================
def registrar_en_bitacora_excel(metadata, id_empresa, giro, tipologia, accion, nombre_archivo="N/A", subcarpeta="N/A"):
    import openpyxl
    from openpyxl.styles import Font, PatternFill
    import datetime as dt_lib
    
    ahora = dt_lib.datetime.now()
    anio_actual = ahora.strftime("%Y")                     
    mes_anio_pestana = ahora.strftime("%B %Y").capitalize() 
    hora_evento = ahora.strftime("%H:%M:%S")
    
    # 🍏 ESCRIBANO UNIVERSAL: HAL localiza la base del Excel al lado de tu script
    dir_script = os.path.dirname(os.path.abspath(__file__))
    carpeta_registros = os.path.join(dir_script, "Registros_PLD")
    if not os.path.exists(carpeta_registros): 
        os.makedirs(carpeta_registros)
    
    ruta_excel_anual = os.path.join(carpeta_registros, f"bitacora_PLD_{anio_actual}.xlsx")

    if os.path.exists(ruta_excel_anual):
        libro = openpyxl.load_workbook(ruta_excel_anual)
    else:
        libro = openpyxl.Workbook()
        libro.remove(libro.active)
    
    titulos = ["Hora", "Analista", "Fecha", "CSV Cargado", "Tipología", "Sujeto", "Giro", "Acción Tomada", "Archivo", "Hipervínculo"]
    
    if mes_anio_pestana in libro.sheetnames:
        hoja = libro[mes_anio_pestana]
    else:
        hoja = libro.create_sheet(title=mes_anio_pestana)
        hoja.append(titulos)
        for col in range(1, len(titulos) + 1):
            hoja.cell(row=1, column=col).font = Font(name="Arial", size=10, bold=True, color="FFFFFF")
            hoja.cell(row=1, column=col).fill = PatternFill(start_color="1E293B", end_color="1E293B", fill_type="solid")
    
    if nombre_archivo != "N/A" and subcarpeta != "N/A":
        formula_enlace = f'=HYPERLINK("../{subcarpeta}/{nombre_archivo}", "🔗 Abrir Expediente")'
    else:
        formula_enlace = "N/A - SIN ARCHIVO FÍSICO"
    
    hoja.append([
        hora_evento, 
        metadata.get("analista", "N/A"), 
        metadata.get("fecha", "N/A"),
        metadata.get("archivo_csv", "N/A"), 
        tipologia, 
        id_empresa, 
        giro, 
        accion, 
        nombre_archivo, 
        formula_enlace
    ])
    
    for col in range(1, hoja.max_column + 1):
        col_letter = openpyxl.utils.get_column_letter(col)
        hoja.column_dimensions[col_letter].width = 18
    
    libro.save(ruta_excel_anual)
    libro.close()


# =====================================================================
# 4. INTERFAZ OPERATIVA DEL ANALISTA (PÁGINA 2 Y 3 UNIVERSAL)
# =====================================================================
# =====================================================================
# 3.5 HOJA 1.5: EL MENÚ MAESTRO DE AVENIDAS OPERATIVAS (BOTONES A Y B)
# =====================================================================
class MenuMaestroAvenidasMitre:
    def __init__(self, root, df, ruta_csv, metadata_auditor):
        self.df = df
        self.meta = metadata_auditor
        self.ruta_csv = ruta_csv
        self.nombre_csv = os.path.basename(ruta_csv)
        self.root = root
        
        self.root.title("Mitre RegTech Core v2026 - Menú Maestro")
        self.root.geometry("640x440")
        self.root.configure(bg="#0f172a")
        
        # Header Superior Ejecutivo
        lbl_top = tk.Label(root, text=f"CONTROL CENTRAL Mitre  |  AUDITOR: {self.meta['analista']}", 
                           font=("Arial", 10, "bold"), fg="#38bdf8", bg="#1e293b", pady=15)
        lbl_top.pack(fill=tk.X)
        
        frame_opciones = tk.LabelFrame(root, text=" 🧭 SELECCIONA LA AVENIDA DE AUDITORÍA FORENSE ", 
                                       font=("Arial", 10, "bold"), fg="#f8fafc", bg="#0f172a", padx=30, pady=25)
        frame_opciones.pack(fill=tk.BOTH, expand=True, padx=40, pady=30)
        
        # 🍏 BOTÓN A: Te lleva directo a tus 5 tipologías CNBV automatizadas que ya funcionan de huevos
        btn_a = tk.Button(frame_opciones, text="🚨 Suite de Alertas Automatizadas (Tipologías CNBV)", 
                          fg="black", bg="#38bdf8", font=("Arial", 10, "bold"), height=2, anchor="w", padx=15,
                          command=self.abrir_avenida_alertas_automatizadas)
        btn_a.pack(fill=tk.X, pady=12)
        
        # 🍏 BOTÓN B: La nueva estación de minería de datos tipo Excel
        btn_b = tk.Button(frame_opciones, text="🔍 Estación Forense de Minería de Datos (Analizar Transacciones)", 
                          fg="black", bg="#059669", font=("Arial", 10, "bold"), height=2, anchor="w", padx=15,
                          command=self.abrir_avenida_mineria_datos_excel)
        btn_b.pack(fill=tk.X, pady=12)
        
                # 4. Colocamos el Riel Inferior de Escape Izquierdo Duplicado
        frame_escape = tk.Frame(frame_opciones, bg="#0f172a")
        frame_escape.pack(fill=tk.X, side=tk.BOTTOM, pady=(25, 0))

        # Botón original de salida definitiva
        btn_cerrar = tk.Button(frame_escape, text=" 🏠 Cerrar Sesión ", fg="black", bg="#94a3b8",
                               font=("Arial", 9, "bold"), bd=0, relief=tk.FLAT, command=self.regresar_a_login)
        btn_cerrar.pack(side=tk.LEFT, padx=5)

    def abrir_avenida_alertas_automatizadas(self):
        # Limpia la ventana de forma segura y manda llamar a tu Hoja 2 original de tipologías
        for widget in self.root.winfo_children():
            widget.destroy()
        SistemaRegTechMitre(self.root, self.df, self.ruta_csv, self.meta)

    def abrir_avenida_mineria_datos_excel(self):
        # 🍏 CONECTOR MITRE: Limpiamos pantalla e invocamos la rejilla tipo Excel usando tu variable real
        for widget in self.root.winfo_children():
            widget.destroy()
        EstacionMineriaDatosMitre(self.root, self.df, self.nombre_csv, self.meta)


    def regresar_a_avenidas(self):
        # Limpiamos el panel de tipologías y despertamos la Hoja 1.5 con tus datos guardados
        for widget in self.root.winfo_children():
            widget.destroy()
        MenuMaestroAvenidasMitre(self.root, self.df, self.nombre_csv, self.meta)

    def regresar_a_menu_maestro(self):
        for widget in self.root.winfo_children(): 
            widget.destroy()
        MenuMaestroAvenidasMitre(self.root, self.df, self.nombre_csv, self.meta)

    def regresar_a_login(self):
        self.root.destroy()
        iniciar_sistema()

class SistemaRegTechMitre:
    def __init__(self, root, df, ruta_csv, metadata_auditor):
        self.df = df
        self.meta = metadata_auditor
        self.nombre_csv = os.path.basename(ruta_csv)
        self.meta["archivo_csv"] = self.nombre_csv
        self.root = root
        self.root.title("Mitre RegTech Core Suite v2026")
        self.root.geometry("780x660")
        self.root.configure(bg="#0f172a")
        
        lbl_top = tk.Label(root, text=f"ESTACIÓN PLD Mitre | AUDITOR: {self.meta['analista']} | BASE: {self.nombre_csv}",
            font=("Arial", 10, "bold"), fg="#38bdf8", bg="#1e293b", pady=12)
        lbl_top.pack(fill=tk.X)
        
        frame_tip = tk.LabelFrame(root, text=" 📊 SELECCIONA TU ENFOQUE DE TIPOLOGÍA CNBV ",
            font=("Arial", 10, "bold"), fg="#f8fafc", bg="#0f172a", padx=15, pady=10)
        frame_tip.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        self.crear_boton_menu(frame_tip, "• Analizar Tipología 1: Velocidad y Fraccionamiento (Pitufeo de Entrada)", "T1")
        self.crear_boton_menu(frame_tip, "• Analizar Tipología 2: Desviación Perfil Histórico (Picos de Actividad)", "T2")
        self.crear_boton_menu(frame_tip, "• Analizar Tipología 3: Inconsistencia de Giro Comercial (Rubro vs Volumen SAT)", "T3")
        self.crear_boton_menu(frame_tip, "• Analizar Tipología 4: Riesgo Geográfico (Cuentas Puente / Islas Caimán)", "T4")
        self.crear_boton_menu(frame_tip, "• Analizar Tipología 5: Umbrales de Ley Planos (Reportes Relevantes ROR)", "T5")
        
        frame_ayuda = tk.LabelFrame(root, text=" 📖 Sustentación Jurídica y Citas de Ley (CNBV México 2026) ",
            font=("Arial", 10, "bold"), fg="#a7f3d0", bg="#1e293b", padx=10, pady=5)
        frame_ayuda.pack(fill=tk.X, padx=20, pady=10)
        
        self.txt_ayuda = tk.Text(frame_ayuda, height=3, font=("Courier", 10), bg="#0f172a", fg="#34d399", relief=tk.FLAT)
        self.txt_ayuda.pack(fill=tk.X)
        self.txt_ayuda.insert(tk.END, ">>> Estación unificada. Selecciona un botón superior para auditar tu lote transaccional transfronterizo...")
        self.txt_ayuda.config(state=tk.DISABLED)
        
        frame_escape = tk.Frame(frame_tip, bg="#0f172a")
        frame_escape.pack(fill=tk.X, side=tk.BOTTOM, pady=(25, 0))
        
        # 🍏 TU LÓGICA MAESTRA: El botón ahora apunta estrictamente a 'self.regresar_a_menu_maestro'
        btn_cerrar = tk.Button(frame_escape, text=" 🔙 Volver al Menú Central ", fg="black", bg="#94a3b8",
                               font=("Arial", 9, "bold"), bd=0, relief=tk.FLAT, command=self.regresar_a_menu_maestro)
        btn_cerrar.pack(side=tk.LEFT, padx=5)

    def crear_boton_menu(self, parent, texto, t_key):
        btn = tk.Button(parent, text=texto, fg="black", bg="#1e293b", font=("Arial", 9, "bold"), anchor="w", padx=15,
            command=lambda: self.mostrar_pantalla_diagnostico_universal(t_key))
        btn.pack(fill=tk.X, pady=4)

    def regresar_a_menu_maestro(self):
            for widget in self.root.winfo_children(): 
                widget.destroy()
            MenuMaestroAvenidasMitre(self.root, self.df, self.nombre_csv, self.meta)

    def regresar_a_avenidas(self):
            # Limpiamos el panel de tipologías y despertamos la Hoja 1.5 con tus datos guardados
            for widget in self.root.winfo_children():
                widget.destroy()
            MenuMaestroAvenidasMitre(self.root, self.df, self.nombre_csv, self.meta)

    def mostrar_pantalla_diagnostico_universal(self, tipo_t):
        for widget in self.root.winfo_children(): 
            widget.destroy()
        
        self.tipologia_actual = tipo_t
        self.datos_alertas_calculadas = {}
        self.empresa_seleccionada_id = None

        tx = {
            "T1": "T1 - VELOCIDAD Y PITUFEO DE ENTRADA",
            "T2": "T2 - DESVIACIÓN PERFIL HISTÓRICO",
            "T3": "T3 - INCONGRUENCIA DE GIRO SAT",
            "T4": "T4 - RIESGO GEOGRÁFICO CUENTAS PUENTE",
            "T5": "T5 - UMBRALES DE LEY RELEVANTES"
        }
        
        lbl_top = tk.Label(self.root, text=tx.get(tipo_t, "AUDITORÍA"), 
                           font=("Arial", 11, "bold"), fg="#f8fafc", bg="#1e293b", pady=12)
        lbl_top.pack(fill=tk.X)

        frame_cuerpo = tk.Frame(self.root, bg="#0f172a")
        frame_cuerpo.pack(fill=tk.BOTH, expand=True, padx=15, pady=2)

        # Columna Izquierda: Tabla y el Mini Resumen
        frame_izq = tk.Frame(frame_cuerpo, bg="#0f172a")
        frame_izq.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))

        self.tabla_alertas = ttk.Treeview(frame_izq, columns=("empresa", "giro", "métrica"), show="headings", height=5)
        self.tabla_alertas.heading("empresa", text="PyME Sospechosa")
        self.tabla_alertas.heading("giro", text="Giro SAT")
        self.tabla_alertas.heading("métrica", text="Métrica Riesgo")
        self.tabla_alertas.column("empresa", width=120, anchor="center")
        self.tabla_alertas.column("giro", width=120, anchor="center")
        self.tabla_alertas.column("métrica", width=110, anchor="center")
        self.tabla_alertas.pack(fill=tk.X, pady=2)

        # Mini Resumen Ejecutivo
        self.txt_sinopsis = tk.Text(frame_izq, height=4, font=("Arial", 10), bg="#1e293b", fg="#e2e8f0", wrap=tk.WORD, relief=tk.FLAT, padx=5, pady=5)
        self.txt_sinopsis.pack(fill=tk.X, pady=2)
        self.txt_sinopsis.insert(tk.END, ">>> Selecciona una PyME para extraer su ADN financiero...")
        self.txt_sinopsis.config(state=tk.DISABLED)

        self.txt_pie_ley = tk.Text(frame_izq, height=2, font=("Courier", 10), bg="#0f172a", fg="#34d399", wrap=tk.WORD, relief=tk.FLAT)
        self.txt_pie_ley.pack(fill=tk.X)
        self.txt_pie_ley.insert(tk.END, ">>> Vinculado a las guías EBR CNBV vigentes en este 2026.")
        self.txt_pie_ley.config(state=tk.DISABLED)

        # Columna Derecha: El lienzo exclusivo de las gráficas vivas
        self.frame_der = tk.LabelFrame(frame_cuerpo, text=" 📊 Firma Geométrica ", font=("Arial", 9, "bold"), fg="#38bdf8", bg="#0f172a")
        self.frame_der.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True)
        
        self.lbl_aviso_grafica = tk.Label(self.frame_der, text="[ Elige una PyME ]", font=("Arial", 9, "italic"), fg="#64748b", bg="#1e293b")
        self.lbl_aviso_grafica.pack(fill=tk.BOTH, expand=True)

        # --- ACCIONES INFERIORES LIMPIAS (ÚNICAMENTE DOS BOTONES) ---
        frame_acc = tk.Frame(self.root, bg="#0f172a", pady=15)
        frame_acc.pack(fill=tk.X, padx=15, side=tk.BOTTOM)
        
        # Botón Izquierdo: Ruta de escape para cancelar o retroceder
        tk.Button(frame_acc, text="🏠 Cancelar / Volver", fg="black", bg="#64748b", 
                  font=("Arial", 10, "bold"), width=20, height=2, 
                  command=self.cancelar_hoja3_cascada).pack(side=tk.LEFT, padx=5)
        
        # Botón Derecho: Disparador directo del Expediente Word/PDF
        tk.Button(frame_acc, text="📝 Generar ROI", fg="black", bg="#38bdf8", 
                  font=("Arial", 10, "bold"), width=28, height=2, 
                  command=self.compilar_reporte_acusacion).pack(side=tk.RIGHT, padx=5)

        # Activadores de telemetría de clics
        self.tabla_alertas.bind("<<TreeviewSelect>>", self.al_seleccionar_fila_universal)
        self.ejecutar_calculos_csv_maestros(tipo_t)
    
    def ejecutar_calculos_csv_maestros(self, tipo):
        if tipo == "T1":
            f = self.df[(self.df['tipo_pago']=='SPEI_RECIBIDO') & (self.df['monto_limpio'].between(9500, 10000))].sort_values(by='fecha_hora')
            res = f.groupby(['id_empresa', 'giro_comercial']).size().reset_index(name='c')
            for _, r in res[res['c']>=5].iterrows():
                emp = str(r['id_empresa']).strip()
                raw = self.df[(self.df['id_empresa']==emp) & (self.df['monto_limpio'].between(9500, 10000))]
                self.datos_alertas_calculadas[emp] = {
                    'expediente': f"ROI-2026-{emp}", 
                    'siglas': 'ROI', 
                    'ventana_tiempo': f"Del {raw['fecha_hora'].min()} al {raw['fecha_hora'].max()}",
                    'cliente_info': [("PyME", emp), ("Giro SAT", r['giro_comercial']), ("Rango", "$9.5k-$10k")],
                    'resumen': f"Fraccionamiento deliberado (Pitufeo) detectada en la cuenta de {emp}. Registró {r['c']} depósitos SPEI.",
                    'movimientos_tabla': [[str(b['fecha_hora']), str(b['tipo_pago']), f"${b['monto_limpio']:,.2f}"] for _, b in raw.iterrows()],
                    'alertas': ["Fraccionamiento bajo umbrales."], 
                    'giro': r['giro_comercial'],
                    'marco_legal': "Violación directa al Artículo 41 de las Disposiciones Generales de la CNBV."
                }
                self.tabla_alertas.insert("", tk.END, values=(emp, r['giro_comercial'], f"{r['c']} SPEIs"))
        
        elif tipo == "T2":
            stats = self.df.groupby('id_empresa').agg(p=('monto_limpio', 'mean'), c=('id_transaccion', 'count')).reset_index()
            df_m = self.df.merge(stats, on='id_empresa')
            for _, row in df_m.iterrows():
                mult = 6 if row['giro_comercial'] in ['Constructora', 'Consultoría TI'] else 3
                if row['c'] >= 3 and row['monto_limpio'] > (row['p'] * mult):
                    emp = str(row['id_empresa']).strip()
                    if emp not in self.datos_alertas_calculadas:
                        raw = self.df[self.df['id_empresa']==emp]
                        self.datos_alertas_calculadas[emp] = {
                            'expediente': f"ROI-T2-{emp}", 
                            'siglas': 'ROI', 
                            'ventana_tiempo': "Periodo de Auditoría 2026",
                            'cliente_info': [("Sujeto", emp), ("Giro", row['giro_comercial']), ("Media Móvil", f"${row['p']:,.2f}")],
                            'resumen': f"Desviación drástica de perfil. El cliente presentó un rascacielos operativo de ${row['monto_limpio']:,.2f} MXN.",
                            'movimientos_tabla': [[str(b['fecha_hora']), str(b['tipo_pago']), f"${b['monto_limpio']:,.2f}"] for _, b in raw.head(5).iterrows()],
                            'alertas': [f"Pico que supera su media."], 
                            'giro': row['giro_comercial'],
                            'marco_legal': "Enfoque Basado en Riesgo CNBV."
                        }
                        self.tabla_alertas.insert("", tk.END, values=(emp, row['giro_comercial'], f"Pico > {mult*100}%"))
        
        elif tipo == "T3":
            cond = (((self.df['giro_comercial'].isin(['Abarrotes', 'Papelería'])) & (self.df['monto_limpio'] > 50000)) | ((self.df['giro_comercial'] == 'Constructora') & (self.df['monto_limpio'] >= 500000)))
            for _, row in self.df[cond].iterrows():
                emp = str(row['id_empresa']).strip()
                if emp not in self.datos_alertas_calculadas:
                    raw = self.df[self.df['id_empresa']==emp]
                    self.datos_alertas_calculadas[emp] = {
                        'expediente': f"ROI-T3-{emp}", 
                        'siglas': 'ROI', 
                        'ventana_tiempo': "Inconsistencia SAT",
                        'cliente_info': [("ID", emp), ("Rubro Declarado", row['giro_comercial']), ("Outlier", f"${row['monto_limpio']:,.2f}")],
                        'resumen': f"Incongruencia de objeto social SAT para {row['giro_comercial']}.",
                        'movimientos_tabla': [[str(b['fecha_hora']), str(b['tipo_pago']), f"${b['monto_limpio']:,.2f}"] for _, b in raw.head(5).iterrows()],
                        'alertas': ["Falta de sustancia económica."], 
                        'giro': row['giro_comercial'],
                        'marco_legal': "Guías de objeto social SAT."
                    }
                    self.tabla_alertas.insert("", tk.END, values=(emp, row['giro_comercial'], f"Outlier: ${row['monto_limpio']:,.2f}"))
        
        elif tipo == "T4":
            for _, row in self.df[(self.df['pais_destino']!='México') & (self.df['tipo_pago']=='SPEI_ENVIADO')].iterrows():
                emp = str(row['id_empresa']).strip()
                if emp not in self.datos_alertas_calculadas:
                    raw = self.df[self.df['id_empresa']==emp]
                    self.datos_alertas_calculadas[emp] = {
                        'expediente': f"ROI-T4-{emp}", 
                        'siglas': 'ROI', 
                        'ventana_tiempo': "Fuga Transfronteriza",
                        'cliente_info': [("PyME Puente", emp), ("Giro", row['giro_comercial']), ("Destino", row['pais_destino'])],
                        'resumen': f"Esquema de Cuenta Puente hacia {row['pais_destino']}.",
                        'movimientos_tabla': [[str(b['fecha_hora']), str(b['tipo_pago']), f"${b['monto_limpio']:,.2f}"] for _, b in raw.head(6).iterrows()],
                        'alertas': ["Fuga internacional offshore."], 
                        'giro': row['giro_comercial'],
                        'marco_legal': "Recomendación 19 del GAFI."
                    }
                    self.tabla_alertas.insert("", tk.END, values=(emp, row['giro_comercial'], row['pais_destino']))
        
        elif tipo == "T5":
            for _, row in self.df[self.df['monto_limpio']>=150000].iterrows():
                emp = str(row['id_empresa']).strip()
                if emp not in self.datos_alertas_calculadas:
                    raw = self.df[self.df['id_empresa']==emp]
                    self.datos_alertas_calculadas[emp] = {
                        'expediente': f"ROR-T5-{emp}", 
                        'siglas': 'ROR', 
                        'ventana_tiempo': "Umbral Plano de Ley",
                        'cliente_info': [("Sujeto ROR", emp), ("Giro", row['giro_comercial']), ("Cruce", f"${row['monto_limpio']:,.2f}")],
                        'resumen': f"Cruce plano regulatorio automatizado de ley.",
                        'movimientos_tabla': [[str(b['fecha_hora']), str(b['tipo_pago']), f"${b['monto_limpio']:,.2f}"] for _, b in raw.head(3).iterrows()],
                        'alertas': ["Cruce de barrera institucional."], 
                        'giro': row['giro_comercial'],
                        'marco_legal': "Artículo 23 de la Ley Fintech."
                    }
                    self.tabla_alertas.insert("", tk.END, values=(emp, row['giro_comercial'], f"${row['monto_limpio']:,.2f}"))
    
    def al_seleccionar_fila_universal(self, event):
        sel = self.tabla_alertas.selection()
        if not sel:
            return
        
        # 🍏 EXTRACCIÓN MAESTRA MAC: Aislamos el ID puro para que haga match perfecto en la memoria del Excel
        valores_fila = self.tabla_alertas.item(sel)['values']
        emp = str(valores_fila[0]).strip() # Captura únicamente el ID limpio de la primera columna
        self.empresa_seleccionada_id = emp
        
        info = self.datos_alertas_calculadas[emp]
        
        self.txt_sinopsis.config(state=tk.NORMAL)
        self.txt_sinopsis.delete("1.0", tk.END)
        self.txt_sinopsis.insert(tk.END, info['resumen'])
        self.txt_sinopsis.config(state=tk.DISABLED)
        
        self.txt_pie_ley.config(state=tk.NORMAL)
        self.txt_pie_ley.delete("1.0", tk.END)
        self.txt_pie_ley.insert(tk.END, f"MARCO: {info['marco_legal']} | Ventana: {info['ventana_tiempo']}")
        self.txt_pie_ley.config(state=tk.DISABLED)
        
        for w in self.frame_der.winfo_children():
            w.destroy()
        
        fig, ax = plt.subplots(figsize=(3.8, 2.8), facecolor="#1e293b")
        ax.set_facecolor("#1e293b")
        ax.tick_params(colors="white", labelsize=7)
        
        raw_m = self.df[self.df['id_empresa'] == emp].sort_values(by='fecha_hora')
        montos = raw_m['monto_limpio'].tolist()
        
        if self.tipologia_actual == "T1":
            ax.bar(range(1, len(montos) + 1), montos, color="crimson", edgecolor="black")
            ax.set_title("Espectro T1: Control Carmesí", color="white", fontsize=8)
        elif self.tipologia_actual == "T2":
            ax.plot(range(1, len(montos) + 1), montos, color="#38bdf8", marker='o', linewidth=1.5)
            ax.set_title("Espectro T2: Rascacielos", color="white", fontsize=8)
        elif self.tipologia_actual == "T3":
            ax.hist([m for m in montos if m < (max(montos) * 0.5)], bins=8, color="#64748b", alpha=0.8)
            ax.axvline(x=max(montos), color="crimson", linewidth=2, label="Outlier")
            ax.set_title("Espectro T3: Objeto Social SAT", color="white", fontsize=8)
        elif self.tipologia_actual == "T4":
            df_rec = self.df['tipo_pago'] == 'SPEI_RECIBIDO'
            df_env = self.df['tipo_pago'] == 'SPEI_ENVIADO'
            df_emp = self.df['id_empresa'] == emp
            rec = self.df[df_emp & df_rec]['monto_limpio'].sum()
            env = self.df[df_emp & df_env]['monto_limpio'].sum()
            ax.bar(['Entrada', 'Fuga'], [rec, env], color=['#64748b', 'crimson'], width=0.4)
            ax.set_title("Espectro T4: Vaciado", color="white", fontsize=8)
        else:
            colores_sc = ['crimson' if m >= 150000 else '#64748b' for m in montos]
            ax.scatter(range(1, len(montos) + 1), montos, color=colores_sc, s=25)
            ax.axhline(y=150000, color="crimson", linestyle="--")
            ax.set_title("Espectro T5: Umbral", color="white", fontsize=8)
        
        plt.tight_layout()
        canvas = FigureCanvasTkAgg(fig, master=self.frame_der)
        canvas.draw()
        canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)
        plt.close(fig)
    
    def cancelar_hoja3_cascada(self):
        if messagebox.askyesno("Navegación", "¿Deseas seguir auditando tipologías en este archivo CSV?"):
            for w in self.root.winfo_children():
                w.destroy()
            self.__init__(self.root, self.df, self.nombre_csv, self.meta)
        else:
            emp = self.empresa_seleccionada_id if self.empresa_seleccionada_id else "N/A"
            registrar_en_bitacora_excel(self.meta, emp, "N/A", self.tipologia_actual, "PROCESO NO COMPLETADO / ANULADO POR EL AUDITOR", "N/A")
            self.root.destroy()
            iniciar_sistema()
    
    def compilar_reporte_acusacion(self):
        if not self.empresa_seleccionada_id:
            messagebox.showerror("Error", "Selecciona una PyME de la tabla.")
            return
            
        emp = self.empresa_seleccionada_id
        datos = self.datos_alertas_calculadas[emp]
        
        # 🍏 CONECTOR DIRECTO PLD: Registramos analista, fecha y el resumen ejecutivo de la PyME
        datos['analista'] = self.meta['analista']
        datos['fecha_auditoria'] = self.meta['fecha']
        datos['entrevista'] = f"Bitácora Forense de Cumplimiento:\n• Análisis de riesgo completado sobre el sujeto {emp}."
        
        n_file = f"{datos['siglas']}_{self.tipologia_actual}_{emp}"
        
        # Registramos la acción en la bitácora anualizada de Excel
        registrar_en_bitacora_excel(
            self.meta, emp, datos['giro'], self.tipologia_actual, 
            f"{datos['siglas']} PREMIUM GENERADO", f"{n_file}.pdf", "Reportes"
        )
        
        # Abrimos el diálogo flotante directo para elegir Word o PDF
        self.ofrecer_empaquetado(
            n_file, "INUSUAL" if self.tipologia_actual != "T5" else "RELEVANTE", datos
        )
    
    def ejecutar_archivo_dictamen_rrd(self):
        if not self.empresa_seleccionada_id:
            messagebox.showerror("Error", "Selecciona una PyME de la tabla.")
            return
        
        emp = self.empresa_seleccionada_id
        dictamen = self.var_dictamen_pld.get()
        notas = self.txt_observaciones_pld.get("1.0", tk.END).strip()
        
        if dictamen == "REPORTAR":
            self.compilar_reporte_acusacion()
            return
        
        txt_m = {
            "DES_LIMPIO": "DESESTIMADO - SIN REVISIÓN", 
            "DES_ESPERA": "DESESTIMADO - EN ESPERA", 
            "DES_3MESES": "DESESTIMADO - LUPA 3 MESES"
        }
        gr_txt = txt_m.get(dictamen, "ARCHIVADO")
        
        datos_c = self.datos_alertas_calculadas[emp]
        datos_rrd = {
            'expediente': f"RRD-2026-{self.tipologia_actual}-{emp}", 
            'grado_txt': gr_txt, 
            'analista': self.meta['analista'], 
            'cliente_info': datos_c['cliente_info'],
            'resumen': f"Se archiva el folio preventivo {self.tipologia_actual} sobre {emp}. Se descarta lavado decretando estatus de {gr_txt}.",
            'observaciones_analista': f"Notas de Cierre:\n{notas if notas else 'Sin notas.'}"
        }
        
        v_r = tk.Toplevel(self.root)
        v_r.title("Exportar RRD")
        v_r.geometry("400x120")
        v_r.configure(bg="#0f172a")
        
        n_f = f"RRD_{self.tipologia_actual}_{emp}"
        
        def g_w():
            crear_word_desestimacion_rrd(f"{n_f}.docx", datos_rrd)
            registrar_en_bitacora_excel(self.meta, emp, datos_c['giro'], self.tipologia_actual, gr_txt, f"{n_f}.docx", "Desestimaciones")
            v_r.destroy()
            messagebox.showinfo("Éxito", "RRD Word guardado.")
            self.cancelar_hoja3_cascada()
        
        def g_p():
            crear_pdf_desestimacion_rrd(f"{n_f}.pdf", datos_rrd)
            registrar_en_bitacora_excel(self.meta, emp, datos_c['giro'], self.tipologia_actual, gr_txt, f"{n_f}.pdf", "Desestimaciones")
            v_r.destroy()
            messagebox.showinfo("Éxito", "RRD PDF guardado.")
            self.cancelar_hoja3_cascada()
        
        tk.Label(v_r, text="¿Formato de guardado para la exoneración RRD?", fg="white", bg="#0f172a", font=("Arial", 9, "bold")).pack(pady=10)
        tk.Button(v_r, text="📝 Descargar Word (.docx)", fg="black", bg="#64748b", font=("Arial", 8, "bold"), width=34, command=g_w).pack(pady=2)
        tk.Button(v_r, text="📕 Descargar PDF (.pdf)", fg="black", bg="#94a3b8", font=("Arial", 8, "bold"), width=34, command=g_p).pack(pady=2)
    
    def ofrecer_empaquetado(self, nombre_base, tipo, datos):
        v = tk.Toplevel(self.root)
        v.geometry("400x130")
        v.configure(bg="#0f172a")
        
        def s_w():
            crear_word_premium(f"{nombre_base}.docx", tipo, datos)
            v.destroy()
            messagebox.showinfo("Éxito", "Word creado.")
            self.cancelar_hoja3_cascada()
        
        def s_p():
            crear_pdf_premium(f"{nombre_base}.pdf", tipo, datos)
            v.destroy()
            messagebox.showinfo("Éxito", "PDF creado.")
            self.cancelar_hoja3_cascada()
        
        tk.Button(v, text="📝 Guardar en Word (.docx)", fg="black", bg="#0284c7", font=("Arial", 9, "bold"), width=36, command=s_w).pack(pady=8)
        tk.Button(v, text="📕 Guardar en PDF (.pdf)", fg="black", bg="#b71c1c", font=("Arial", 9, "bold"), width=36, command=s_p).pack(pady=4)

# =====================================================================
# 3.8 ESTACIÓN FORENSE DE MINERÍA DE DATOS (REJILLA COMPLETA INTERACTIVA)
# =====================================================================
class EstacionMineriaDatosMitre:
    def __init__(self, root, df, nombre_csv, metadata_auditor):
        self.df = df
        self.meta = metadata_auditor
        self.nombre_csv = nombre_csv
        self.root = root
        
        self.root.title("Mitre RegTech Core - Explorador Forense de Transacciones")
        self.root.geometry("920x680") # Ampliamos un poco el ancho para los nuevos filtros
        self.root.configure(bg="#0f172a")
        
        # Header Superior Corporativo Mitre
        lbl_top = tk.Label(root, text=f"MITRE PLD ANALYTICS  |  EXPLORADOR DE DATA CRUDA  |  BASE: {nombre_csv}", 
                           font=("Arial", 10, "bold"), fg="#38bdf8", bg="#1e293b", pady=12)
        lbl_top.pack(fill=tk.X)
        
        # 🔍 BARDA SUPERIOR REFORZADA DE FILTROS (DISEÑO GRID ORDENADO)
        frame_filtros = tk.LabelFrame(root, text=" 🔍 Filtros Forenses Avanzados (Motores Pandas) ", 
                                       font=("Arial", 9, "bold"), fg="#a7f3d0", bg="#0f172a", padx=10, pady=10)
        frame_filtros.pack(fill=tk.X, padx=15, pady=8)
        
        # FILA 0: Filtros de Identidad y Dinero
        tk.Label(frame_filtros, text="ID PyME:", fg="#94a3b8", bg="#0f172a", font=("Arial", 8, "bold")).grid(row=0, column=0, padx=5, pady=2, sticky="w")
        self.ent_filter_emp = tk.Entry(frame_filtros, width=12, font=("Arial", 9), bg="#1e293b", fg="white", insertbackground="white")
        self.ent_filter_emp.grid(row=0, column=1, padx=5, pady=2)
        
        tk.Label(frame_filtros, text="Monto Mín:", fg="#94a3b8", bg="#0f172a", font=("Arial", 8, "bold")).grid(row=0, column=2, padx=5, pady=2, sticky="w")
        self.ent_filter_min = tk.Entry(frame_filtros, width=10, font=("Arial", 9), bg="#1e293b", fg="white", insertbackground="white")
        self.ent_filter_min.grid(row=0, column=3, padx=5, pady=2)
        
        tk.Label(frame_filtros, text="Giro SAT:", fg="#94a3b8", bg="#0f172a", font=("Arial", 8, "bold")).grid(row=0, column=4, padx=5, pady=2, sticky="w")
        self.ent_filter_giro = tk.Entry(frame_filtros, width=15, font=("Arial", 9), bg="#1e293b", fg="white", insertbackground="white")
        self.ent_filter_giro.grid(row=0, column=5, padx=5, pady=2)

        # FILA 1: Filtros de Mecanismo y Geografía
        tk.Label(frame_filtros, text="Tipo SPEI:", fg="#94a3b8", bg="#0f172a", font=("Arial", 8, "bold")).grid(row=1, column=0, padx=5, pady=4, sticky="w")
        self.combo_filter_pago = ttk.Combobox(frame_filtros, values=["TODOS", "SPEI_RECIBIDO", "SPEI_ENVIADO"], width=13, state="readonly")
        self.combo_filter_pago.set("TODOS")
        self.combo_filter_pago.grid(row=1, column=1, padx=5, pady=4)
        
        tk.Label(frame_filtros, text="País Destino:", fg="#94a3b8", bg="#0f172a", font=("Arial", 8, "bold")).grid(row=1, column=2, padx=5, pady=4, sticky="w")
        self.ent_filter_pais = tk.Entry(frame_filtros, width=12, font=("Arial", 9), bg="#1e293b", fg="white", insertbackground="white")
        self.ent_filter_pais.grid(row=1, column=3, padx=5, pady=4)
        
        # Botón de Disparo Ejecutivo (Gatillo Pandas)
        btn_run = tk.Button(frame_filtros, text="⚡ Ejecutar Minería", fg="black", bg="#34d399", 
                            font=("Arial", 9, "bold"), width=16, command=self.ejecutar_filtrado_pandas)
        btn_run.grid(row=1, column=5, padx=15, pady=4, columnspan=2, sticky="e")

        # Rejilla Visual Estilo Excel
        frame_tabla = tk.Frame(root, bg="#0f172a")
        frame_tabla.pack(fill=tk.BOTH, expand=True, padx=15, pady=5)
        
        columnas = ("fecha", "empresa", "giro", "pago", "monto", "destino")
        self.tabla_excel = ttk.Treeview(frame_tabla, columns=columnas, show="headings", height=16)
        
        self.tabla_excel.heading("fecha", text="Fecha / Hora")
        self.tabla_excel.heading("empresa", text="ID Empresa")
        self.tabla_excel.heading("giro", text="Giro SAT")
        self.tabla_excel.heading("pago", text="Tipo Pago")
        self.tabla_excel.heading("monto", text="Monto (MXN)")
        self.tabla_excel.heading("destino", text="País Destino")
        
        for col in columnas:
            self.tabla_excel.column(col, width=140, anchor="center")
            
        scroll_v = ttk.Scrollbar(frame_tabla, orient=tk.VERTICAL, command=self.tabla_excel.yview)
        self.tabla_excel.configure(yscrollcommand=scroll_v.set)
        
        self.tabla_excel.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scroll_v.pack(side=tk.RIGHT, fill=tk.Y)

        # Riel de Acciones Inferior
        frame_bot = tk.Frame(root, bg="#0f172a", pady=10)
        frame_bot.pack(fill=tk.X, padx=15, side=tk.BOTTOM)
        
        tk.Button(frame_bot, text="🔙 Volver al Menú Central", fg="black", bg="#94a3b8", 
                  font=("Arial", 9, "bold"), command=self.regresar_a_menu_maestro).pack(side=tk.LEFT)
                  
        self.lbl_contador = tk.Label(frame_bot, text="Registros cargados: 0", font=("Arial", 9, "bold"), fg="#38bdf8", bg="#0f172a")
        self.lbl_contador.pack(side=tk.RIGHT)
        
        self.ejecutar_filtrado_pandas()

    def ejecutar_filtrado_pandas(self):
        for item in self.tabla_excel.get_children():
            self.tabla_excel.delete(item)
            
        df_temp = self.df.copy()
        
        # 🍏 1. Filtro ID PyME
        emp_val = self.ent_filter_emp.get().strip()
        if emp_val:
            df_temp = df_temp[df_temp['id_empresa'].astype(str).str.contains(emp_val, case=False)]
            
        # 🍏 2. Filtro Monto Mínimo
        min_val = self.ent_filter_min.get().strip()
        if min_val:
            try: df_temp = df_temp[df_temp['monto_limpio'] >= float(min_val)]
            except: pass
            
        # 🍏 3. Filtro Giro SAT (Nuevo)
        giro_val = self.ent_filter_giro.get().strip()
        if giro_val:
            df_temp = df_temp[df_temp['giro_comercial'].astype(str).str.contains(giro_val, case=False)]
            
        # 🍏 4. Filtro Tipo SPEI (Nuevo Menú Desplegable)
        pago_val = self.combo_filter_pago.get()
        if pago_val != "TODOS":
            df_temp = df_temp[df_temp['tipo_pago'] == pago_val]
            
        # 🍏 5. Filtro País Destino (Nuevo)
        pais_val = self.ent_filter_pais.get().strip()
        if pais_val:
            df_temp = df_temp[df_temp['pais_destino'].astype(str).str.contains(pais_val, case=False)]
            
        # Surtimos la rejilla Excel con la data filtrada de Pandas
        for _, r in df_temp.iterrows():
            monto_f = f"${r['monto_limpio']:,.2f}"
            self.tabla_excel.insert("", tk.END, values=(str(r['fecha_hora']), str(r['id_empresa']), str(r['giro_comercial']), str(r['tipo_pago']), monto_f, str(r['pais_destino'])))
            
        self.lbl_contador.config(text=f"Registros Encontrados: {len(df_temp)}")

    def regresar_a_menu_maestro(self):
        for widget in self.root.winfo_children(): 
            widget.destroy()
        # Llamamos a tu clase corregida Mitre de forma perfecta
        MenuMaestroAvenidasMitre(self.root, self.df, self.nombre_csv, self.meta)

# =====================================================================
# 5. SUITE EXCLUSIVA DEL SUPERADMINISTRADOR (HOJA 2 Y 3 ADMIN)
# =====================================================================
class PanelAdminMitre:
    def __init__(self, root):
        self.root = root
        self.root.title("Mitre RegTech Core - Consola del Administrador")
        self.root.geometry("600x400")
        self.root.configure(bg="#0f172a")
        
        lbl_top = tk.Label(root, text="CONSOLA DE CONFIGURACIÓN Y GOBIERNO CORPORATIVO\n[ ACCESO EXCLUSIVO ADM PLD 2026 ]",
            font=("Arial", 11, "bold"), fg="#a7f3d0", bg="#1e293b", pady=15)
        lbl_top.pack(fill=tk.X)
        
        frame_avenidas = tk.Frame(root, bg="#0f172a", pady=35)
        frame_avenidas.pack(fill=tk.BOTH, expand=True, padx=40)
        
        tk.Button(frame_avenidas, text="📂 AUDITAR EXPEDIENTES EMITIDOS (Word / PDF)", fg="black", bg="#0284c7", font=("Arial", 10, "bold"), height=2, command=self.abrir_hoja_3_admin_reportes).pack(fill=tk.X, pady=10)
        tk.Button(frame_avenidas, text="📊 REVISAR BITÁCORAS ANUALES DE ACTIVIDAD (Excel)", fg="black", bg="#059669", font=("Arial", 10, "bold"), height=2, command=self.abrir_hoja_3_admin_registros).pack(fill=tk.X, pady=10)
        
        frame_escape = tk.Frame(root, bg="#0f172a")
        frame_escape.pack(fill=tk.X, side=tk.BOTTOM, pady=(0, 20), padx=40)
        tk.Button(frame_escape, text=" 🏠 Cerrar Sesión / Volver al Login ", fg="black", bg="#94a3b8", font=("Arial", 9, "bold"), bd=0, command=self.regresar_a_pagina_1_admin).pack(side=tk.LEFT)
    
    def regresar_a_pagina_1_admin(self):
        self.root.destroy()
        lanzar_pantalla_registro_oficial()
    
    def abrir_hoja_3_admin_reportes(self):
        for widget in self.root.winfo_children(): 
            widget.destroy()
        tk.Label(self.root, text="CONSOLA ADMINISTRATIVA: AUDITORÍA DE EXPEDIENTES", font=("Arial", 11, "bold"), fg="#38bdf8", bg="#1e293b", pady=12).pack(fill=tk.X)
        
        frame_filtros = tk.LabelFrame(self.root, text=" 🔍 Filtros de Segmentación ", font=("Arial", 9, "bold"), fg="#a7f3d0", bg="#0f172a", padx=10, pady=10)
        frame_filtros.pack(fill=tk.X, padx=15, pady=10)
        self.combo_mes = ttk.Combobox(frame_filtros, values=["Enero", "Febrero", "Marzo", "Abril", "Mayo", "Junio", "Julio", "Agosto", "Septiembre", "Octubre", "Noviembre", "Diciembre"], width=12, state="readonly")
        self.combo_mes.set("Agosto")
        self.combo_mes.grid(row=0, column=0, padx=10)
        tk.Button(frame_filtros, text="⚡ Escanear Finder Mac", fg="black", bg="#34d399", font=("Arial", 8, "bold"), command=self.scan_f).grid(row=0, column=1, padx=10)

        self.tabla_archivos_admin = ttk.Treeview(self.root, columns=("archivo", "tipo"), show="headings", height=8)
        self.tabla_archivos_admin.heading("archivo", text="Expediente Detectado")
        self.tabla_archivos_admin.heading("tipo", text="Tipo")
        self.tabla_archivos_admin.column("archivo", width=340)
        self.tabla_archivos_admin.column("tipo", width=120)
        self.tabla_archivos_admin.pack(fill=tk.BOTH, expand=True, padx=15)
        self.scan_f()

        frame_acc = tk.Frame(self.root, bg="#0f172a", pady=15)
        frame_acc.pack(fill=tk.X, padx=15, side=tk.BOTTOM)
        tk.Button(frame_acc, text="🔙 Volver al Panel Admin", fg="black", bg="#94a3b8", font=("Arial", 9, "bold"), width=20, command=self.reload_menu).pack(side=tk.LEFT)
        tk.Button(frame_acc, text="🔗 Abrir Expediente Oficial", fg="black", bg="#0284c7", font=("Arial", 9, "bold"), width=24, command=self.open_f).pack(side=tk.RIGHT)

    def scan_f(self):
        # 1. Limpiamos la tabla por completo
        for item in self.tabla_archivos_admin.get_children(): 
            self.tabla_archivos_admin.delete(item)
            
        # 2. Mapeamos el mes seleccionado a su número correspondiente
        meses_dict = {
            "ENERO": 1, "FEBRERO": 2, "MARZO": 3, "ABRIL": 4, 
            "MAYO": 5, "JUNIO": 6, "JULIO": 7, "AGOSTO": 8, 
            "SEPTIEMBRE": 9, "OCTUBRE": 10, "NOVIEMBRE": 11, "DICIEMBRE": 12
        }
        mes_elegido = self.combo_mes.get().upper()
        numero_mes_esperado = meses_dict.get(mes_elegido, 8)
        
        c = os.path.join(os.getcwd(), "Reportes")
        if os.path.exists(c):
            import datetime
            # 3. Escaneamos tu Finder físico
            for f in os.listdir(c):
                if f.startswith("."): 
                    continue # Ignoramos basura oculta del sistema Mac
                
                ruta_archivo = os.path.join(c, f)
                
                # 🍏 TRUCO DE TELEMETRÍA: HAL lee cuándo se modificó el archivo de verdad
                timestamp = os.path.getmtime(ruta_archivo)
                fecha_modificacion = datetime.datetime.fromtimestamp(timestamp)
                
                # Si el mes en que se guardó coincide con el que buscas, brota en la tabla
                if fecha_modificacion.month == numero_mes_esperado:
                    self.tabla_archivos_admin.insert("", tk.END, values=(f, "Reporte PLD"))

    def open_f(self):
        sel = self.tabla_archivos_admin.selection()
        if sel:
            f_vals = self.tabla_archivos_admin.item(sel)['values']
            nombre_archivo = f_vals[0] if isinstance(f_vals, list) else str(f_vals)
            os.system(f'open "{os.path.join(os.getcwd(), "Reportes", nombre_archivo)}"')

    def abrir_hoja_3_admin_registros(self):
        for widget in self.root.winfo_children(): 
            widget.destroy()
            
        tk.Label(self.root, text="CONSOLA ADMINISTRATIVA: BITÁCORAS DE ACTIVIDAD (EXCEL)", 
                 font=("Arial", 11, "bold"), fg="#059669", bg="#1e293b", pady=12).pack(fill=tk.X)
        
        # 🔍 Creamos el contenedor elegante de filtros para el año
        frame_filtros = tk.LabelFrame(self.root, text=" 🔍 Filtro por Año de Actividad ", font=("Arial", 9, "bold"), fg="#a7f3d0", bg="#0f172a", padx=10, pady=10)
        frame_filtros.pack(fill=tk.X, padx=15, pady=10)
        
        # Menú desplegable de años restringido (Read-Only)
        self.combo_anio = ttk.Combobox(frame_filtros, values=["2024", "2025", "2026", "2027"], width=12, state="readonly")
        self.combo_anio.set("2026")
        self.combo_anio.grid(row=0, column=0, padx=10)
        
        # El botón verde ahora invoca a la nueva función inteligente 'scan_x'
        tk.Button(frame_filtros, text="⚡ Escanear Libros Excel", fg="black", bg="#34d399", font=("Arial", 8, "bold"), command=self.scan_x).grid(row=0, column=1, padx=10)

        self.tabla_excels_admin = ttk.Treeview(self.root, columns=("archivo", "tipo"), show="headings", height=8)
        self.tabla_excels_admin.heading("archivo", text="Libro de Actividades Excel Anual (.xlsx)")
        self.tabla_excels_admin.heading("tipo", text="Formato")
        self.tabla_excels_admin.pack(fill=tk.BOTH, expand=True, padx=15, pady=10)

        frame_acc = tk.Frame(self.root, bg="#0f172a", pady=15)
        frame_acc.pack(fill=tk.X, padx=15, side=tk.BOTTOM)
        tk.Button(frame_acc, text="🔙 Volver al Panel Admin", fg="black", bg="#94a3b8", font=("Arial", 9, "bold"), width=20, command=self.reload_menu).pack(side=tk.LEFT)
        tk.Button(frame_acc, text="📊 Abrir Bitácora Completa en Excel", fg="black", bg="#059669", font=("Arial", 9, "bold"), width=28, command=self.open_x).pack(side=tk.RIGHT)

    def scan_x(self):
        # 1. Limpiamos la tabla de Excels por completo
        for item in self.tabla_excels_admin.get_children(): 
            self.tabla_excels_admin.delete(item)
            
        # 2. 🍏 CONECTOR DE AÑO: Absorbemos el año seleccionado en la pantalla
        anio_elegido = self.combo_anio.get().strip() # Ej. "2026"
        
        c = os.path.join(os.getcwd(), "Registros_PLD")
        if os.path.exists(c):
            # 3. Escaneamos la carpeta física en tu Mac
            for f in os.listdir(c):
                if f.endswith(".xlsx") and not f.startswith("."):
                    f_upper = f.upper()
                    
                    # 🍏 EL FILTRO INTERACTIVO: Si el nombre del Excel contiene el año buscado
                    if anio_elegido in f_upper:
                        self.tabla_excels_admin.insert("", tk.END, values=(f, "Libro Excel Anual"))

    def open_x(self):
        sel = self.tabla_excels_admin.selection()
        if sel:
            f_vals = self.tabla_excels_admin.item(sel)['values']
            nombre_excel = f_vals[0] if isinstance(f_vals, list) else str(f_vals)
            os.system(f'open "{os.path.join(os.getcwd(), "Registros_PLD", nombre_excel)}"')

    def reload_menu(self):
        for w in self.root.winfo_children(): 
            w.destroy()
        self.__init__(self.root)

# =====================================================================
# 6. PANTALLA INICIAL DE ACCESO DE SEGURIDAD (LOGIN - HOJA 1)
# =====================================================================
def lanzar_pantalla_registro_oficial():
    login_win = tk.Tk()
    login_win.title("Registro de Control Mitre PLD")
    login_win.geometry("400x310")
    login_win.configure(bg="#0f172a")
    
    tk.Label(login_win, text="ACCESO RESTRINGIDO: INGRESO DE AUDITOR", font=("Arial", 11, "bold"), fg="#38bdf8", bg="#0f172a", pady=15).pack()
    
    tk.Label(login_win, text="Nombre Completo del Analista Responsable:", fg="#94a3b8", bg="#0f172a", font=("Arial", 9, "bold")).pack(anchor="w", padx=40, pady=2)
    ent_analista = tk.Entry(login_win, width=32, font=("Arial", 10), bg="#1e293b", fg="white", insertbackground="white")
    ent_analista.pack(padx=40, pady=2)
    ent_analista.insert(0, "Lic. Ana Torres Ramírez")
    
    tk.Label(login_win, text="Fecha de la Investigación (AAAA-MM-DD):", fg="#94a3b8", bg="#0f172a", font=("Arial", 9, "bold")).pack(anchor="w", padx=40, pady=2)
    ent_fecha = tk.Entry(login_win, width=32, font=("Arial", 10), bg="#1e293b", fg="white", insertbackground="white")
    ent_fecha.pack(padx=40, pady=2)
    
    import datetime as dt_actual
    fecha_hoy_pld = dt_actual.datetime.now().strftime("%Y-%m-%d")
    ent_fecha.insert(0, fecha_hoy_pld)
    
    def registrar_y_continuar():
        analista, fecha = ent_analista.get().strip(), ent_fecha.get().strip()
        if not analista or not fecha:
            messagebox.showerror("Error", "Campos obligatorios.")
            return
        
        metadata = {"analista": analista, "fecha": fecha}
        login_win.destroy()
        
        root_f = tk.Tk()
        root_f.withdraw()
        messagebox.showinfo("Carga Core Mitre", "Registro completado. Elige tu archivo CSV transaccional.")
        
        ruta_csv = filedialog.askopenfilename(title="Seleccionar CSV", filetypes=[("CSV", "*.csv")])
        if ruta_csv:
            df = pd.read_csv(ruta_csv)
            df.columns = df.columns.str.lower().str.strip()
            df['monto_limpio'] = df['monto'].astype(str).str.replace(r'[^0-9.]', '', regex=True).str.strip().astype(float)
                        # 🍏 AJUSTE DIRECTO: Jala el texto original completo (Fecha + Hora) sin recortar nada
            df['fecha_hora'] = df['fecha_hora'].astype(str).str.strip()
            
            # 🍏 ENLACE MAESTRO: HAL ahora arranca en la Hoja 1.5 del Menú de Opciones A y B
            root = tk.Tk()
            MenuMaestroAvenidasMitre(root, df, ruta_csv, metadata)
            root.mainloop()

    
    tk.Button(login_win, text="🔒 Validar Identidad y Cargar CSV", fg="black", bg="#059669", font=("Arial", 10, "bold"), width=24, height=2, command=registrar_y_continuar).pack(pady=15)
    
    tk.Label(login_win, text="________________________________________", fg="#334155", bg="#0f172a").pack()
    
    def ir_admin():
        if simpledialog.askstring("Área Restringida", "Ingresa la Clave de Superadministrador:", show="*") == "1234":
            login_win.destroy()
            r_a = tk.Tk()
            PanelAdminMitre(r_a)
            r_a.mainloop()
        else:
            messagebox.showerror("Error", "Denegado.")
    
    tk.Button(login_win, text="⚙️ Consola del Administrador", fg="black", bg="#475569", font=("Arial", 8, "bold"), bd=0, command=ir_admin).pack(side=tk.LEFT, padx=20, pady=10)
    
    login_win.mainloop()


def iniciar_sistema():
    lanzar_pantalla_registro_oficial()


if __name__ == "__main__":
    iniciar_sistema()
